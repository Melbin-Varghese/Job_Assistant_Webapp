"""
ATS Resume Improver — Flask blueprint
------------------------------------
Single-page job-assistance web app.

Flow:
  1. User uploads a resume (.pdf / .docx) and pastes a job description.
  2. We extract resume text, extract JD keywords, compute an ATS match score.
  3. We parse the resume into structured sections and weave in the missing
     JD keywords (mainly into the Skills section) to produce an improved version.
  4. We render that improved resume as a single-page, ATS-safe PDF
     (no tables/columns/images — just clean text a parser can read),
     shrinking font size / trimming bullets as needed to keep it to one page.
  5. User downloads the tailored PDF.
"""

import os
import re
import uuid
import string
import tempfile
from collections import Counter

from flask import Blueprint, render_template, request, jsonify, send_file

import docx  # python-docx
from pypdf import PdfReader

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.pdfbase.pdfmetrics import stringWidth

# --------------------------------------------------------------------------
# Blueprint setup
# --------------------------------------------------------------------------

resume_bp = Blueprint('resume', __name__, url_prefix='/resume')

# Use temp directory for file storage
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {"pdf", "docx"}

STOPWORDS = set("""
a about above after again against all am an and any are aren't as at be because
been before being below between both but by can't cannot could couldn't did
didn't do does doesn't doing don't down during each few for from further had
hadn't has hasn't have haven't having he he'd he'll he's her here here's hers
herself him himself his how how's i i'd i'll i'm i've if in into is isn't it
it's its itself let's me more most mustn't my myself no nor not of off on once
only or other ought our ours ourselves out over own same shan't she she'd
she'll she's should shouldn't so some such than that that's the their theirs
them themselves then there there's these they they'd they'll they're they've
this those through to too under until up very was wasn't we we'd we'll we're
we've were weren't what what's when when's where where's which while who
who's whom why why's with won't would wouldn't you you'd you'll you're you've
your yours yourself yourselves using use used work experience years strong
excellent good etc will able within team role responsibilities requirements
preferred required plus job company looking seeking ideal candidate join
including includes include etc must nice ability skills skill knowledge
skilled proficient familiarity familiar demonstrated
""".split())

SECTION_HEADERS = {
    "summary": ["SUMMARY", "PROFILE", "OBJECTIVE", "PROFESSIONAL SUMMARY", "CAREER SUMMARY"],
    "skills": ["SKILLS", "CORE SKILLS", "TECHNICAL SKILLS", "KEY SKILLS", "SKILLS & TOOLS"],
    "experience": ["EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE",
                   "EMPLOYMENT HISTORY", "WORK HISTORY"],
    "education": ["EDUCATION", "EDUCATION & TRAINING", "ACADEMIC BACKGROUND"],
    "projects": ["PROJECTS", "KEY PROJECTS", "PERSONAL PROJECTS"],
    "certifications": ["CERTIFICATIONS", "CERTIFICATES", "LICENSES"],
    "achievements": ["ACHIEVEMENTS", "AWARDS", "ACCOMPLISHMENTS"],
}
ALL_HEADER_STRINGS = {h: k for k, arr in SECTION_HEADERS.items() for h in arr}

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d\-\s()]{7,}\d)")
LINK_RE = re.compile(r"(https?://\S+|linkedin\.com/\S+|github\.com/\S+)", re.I)

# PDF constants
PAGE_W, PAGE_H = LETTER
MARGIN = 0.5 * inch
CONTENT_W = PAGE_W - 2 * MARGIN
FONT_HEADER = "Helvetica-Bold"
FONT_BOLD = "Helvetica-Bold"
FONT_BODY = "Helvetica"


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------

def extract_text_from_pdf(path):
    text = []
    with open(path, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            text.append(page.extract_text() or "")
    return "\n".join(text)


def extract_text_from_docx(path):
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_text(path, filename):
    ext = filename.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(path)
    elif ext == "docx":
        return extract_text_from_docx(path)
    return ""


# --------------------------------------------------------------------------
# Keyword / ATS scoring (job description vs resume)
# --------------------------------------------------------------------------

def tokenize(text):
    text = text.lower()
    text = re.sub(r"[^a-z0-9\s\+\#\.\-/]", " ", text)
    tokens = text.split()
    cleaned = []
    for t in tokens:
        t = t.strip(string.punctuation)
        t = t.strip("-/.")
        if len(t) < 2 or t in STOPWORDS or t.isdigit():
            continue
        cleaned.append(t)
    return cleaned


def extract_keywords(jd_text, top_n=40):
    tokens = tokenize(jd_text)
    freq = Counter(tokens)

    words = jd_text.lower().split()
    bigrams = []
    for i in range(len(words) - 1):
        w1 = re.sub(r"[^a-z0-9\+\#]", "", words[i])
        w2 = re.sub(r"[^a-z0-9\+\#]", "", words[i + 1])
        if w1 in STOPWORDS or w2 in STOPWORDS or len(w1) < 2 or len(w2) < 2:
            continue
        bigrams.append(f"{w1} {w2}")
    bigram_freq = Counter(bigrams)

    keywords = [w for w, _ in freq.most_common(top_n)]
    top_bigrams = [b for b, c in bigram_freq.most_common(15) if c > 1]
    return keywords, top_bigrams


def compute_match(resume_text, jd_keywords, jd_bigrams):
    resume_tokens = set(tokenize(resume_text))
    resume_lower = resume_text.lower()

    matched = [kw for kw in jd_keywords if kw in resume_tokens]
    missing = [kw for kw in jd_keywords if kw not in resume_tokens]

    matched_bigrams = [b for b in jd_bigrams if b in resume_lower]
    missing_bigrams = [b for b in jd_bigrams if b not in resume_lower]

    total = len(jd_keywords) + len(jd_bigrams)
    hit = len(matched) + len(matched_bigrams)
    score = round((hit / total) * 100, 1) if total else 0.0

    return {
        "score": score,
        "matched": matched,
        "missing": missing,
        "matched_bigrams": matched_bigrams,
        "missing_bigrams": missing_bigrams,
    }


# --------------------------------------------------------------------------
# Resume parsing into structured sections
# --------------------------------------------------------------------------

def detect_header(line):
    """Return the canonical section key if this line looks like a section header."""
    stripped = line.strip().strip(":").strip()
    if not stripped or len(stripped.split()) > 5:
        return None
    upper = stripped.upper()
    if upper in ALL_HEADER_STRINGS:
        return ALL_HEADER_STRINGS[upper]
    # Fallback: an ALL-CAPS short line not otherwise matched is still a header
    if stripped.isupper() and 1 <= len(stripped.split()) <= 5:
        return "other:" + stripped.title()
    return None


def parse_resume(text):
    raw_lines = [l.strip() for l in text.splitlines()]
    lines = [l for l in raw_lines if l != ""]

    # Extract header (first few lines before first section)
    header_lines = []
    header_end = 0
    for i, line in enumerate(lines[:10]):
        if detect_header(line):
            header_end = i
            break
        header_lines.append(line)

    parsed = {
        "header": header_lines,
        "summary": [],
        "skills": [],
        "experience": [],
        "education": [],
        "projects": [],
        "certifications": [],
        "achievements": [],
        "other": [],
    }

    current_section = None
    current_entry = None
    lines_to_process = lines[header_end:] if header_end > 0 else lines

    for line in lines_to_process:
        header = detect_header(line)
        if header and header not in parsed:
            parsed[header] = []
            current_section = header
            current_entry = None
            continue

        if header:
            if current_entry and current_section in ("experience", "education", "projects", "other"):
                parsed[current_section].append(current_entry)
                current_entry = None
            current_section = header
            continue

        if not current_section:
            continue

        section_type = current_section if not current_section.startswith("other:") else "other"

        if section_type == "summary":
            parsed[section_type].append(line)
        elif section_type == "skills":
            for skill in re.split(r"[,•]", line):
                skill = skill.strip()
                if skill:
                    parsed[section_type].append(skill)
        elif section_type in ("experience", "education", "projects", "other"):
            if line.startswith("•") or line.startswith("-"):
                bullet = line.lstrip("•-").strip()
                if current_entry:
                    current_entry["bullets"].append(bullet)
            else:
                if current_entry and "bullets" in current_entry:
                    parsed[section_type].append(current_entry)
                current_entry = {
                    "heading": line,
                    "bullets": []
                }
        elif section_type in ("certifications", "achievements"):
            bullet = line.lstrip("•-").strip() if line.startswith(("•", "-")) else line
            if bullet:
                parsed[section_type].append(bullet)

    if current_entry and current_section in ("experience", "education", "projects", "other"):
        parsed[current_section].append(current_entry)

    return parsed


def improve_resume(parsed, missing_keywords, missing_bigrams, jd_text="", job_role=""):
    improved = {k: v[:] for k, v in parsed.items()}
    added = []

    # Clean up skills - remove junk/short words and duplicates
    improved["skills"] = list(dict.fromkeys([s for s in improved["skills"] 
                                             if len(s) > 3 and s.lower() not in STOPWORDS]))

    # Add missing keywords to skills (top 10 only)
    to_add = list(missing_keywords) + list(missing_bigrams)
    to_add = list(dict.fromkeys(to_add))[:10]

    if to_add:
        improved["skills"].extend(to_add)
        added = to_add

    # Generate professional summary from JD keywords and job role
    if not improved["summary"] or (improved["summary"] and len(" ".join(improved["summary"])) < 100):
        professional_summary = generate_professional_summary(missing_keywords, missing_bigrams, job_role)
        improved["summary"] = [professional_summary] if professional_summary else improved["summary"]

    improved["_added_skills"] = added
    return improved


def generate_professional_summary(keywords, bigrams, job_role=""):
    """Generate a professional 3-4 line summary using top JD keywords and job role."""
    if not keywords and not job_role:
        return ""
    
    # Use top keywords
    top_keywords = keywords[:6]
    key_bigrams = [b for b in bigrams[:2] if b]
    
    # Build summary - 3-4 lines
    if job_role:
        summary = f"Results-driven {job_role} with expertise in {', '.join(top_keywords[:3])}. "
    else:
        summary = f"Results-driven professional with expertise in {', '.join(top_keywords[:3])}. "
    
    summary += f"Skilled in {', '.join(top_keywords[3:6])}. "
    
    if key_bigrams:
        summary += f"Specialized in {', '.join(key_bigrams)}. "
    
    summary += f"Seeking to deliver innovative solutions in {job_role if job_role else 'technology'}."
    
    return summary.strip()


# --------------------------------------------------------------------------
# PDF rendering
# --------------------------------------------------------------------------

def wrap_text(text, font, size, width):
    """Wrap text to fit within a given width, accounting for proportional fonts."""
    lines = []
    words = text.split()
    current_line = ""

    for word in words:
        test_line = (current_line + " " + word).strip()
        w = stringWidth(test_line, font, size)
        if w <= width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = word

    if current_line:
        lines.append(current_line)
    return lines


def build_pdf_lines(improved, body_size, max_bullets_per_entry, max_summary):
    items = []

    def section_title(title):
        items.append((title, FONT_HEADER, body_size + 1.5, 8, 0))
        items.append(("__RULE__", None, 0, 3, 0))

    gap_entry = 6

    # Add professional header section (name, contact info)
    if improved.get("header"):
        # Extract name (usually first line)
        name_line = improved["header"][0] if improved["header"] else ""
        contact_lines = improved["header"][1:] if len(improved["header"]) > 1 else []
        
        if name_line.strip():
            # Add name in bold, larger font
            items.append((name_line.strip(), FONT_BOLD, body_size + 2, 0, 0))
        
        # Add contact info lines (smaller, centered feel)
        for contact in contact_lines[:3]:  # Limit to 3 contact lines
            if contact.strip():
                items.append((contact.strip(), FONT_BODY, body_size - 1, 2, 0))
        
        items.append(("__RULE__", None, 0, 6, 0))

    # Professional Summary section
    if improved["summary"]:
        section_title("PROFESSIONAL SUMMARY")
        for s in improved["summary"][:max_summary]:
            if s.strip():
                for l in wrap_text(s.strip(), FONT_BODY, body_size, CONTENT_W):
                    items.append((l, FONT_BODY, body_size, 0, 0))

    # Skills section - ATS format
    if improved["skills"]:
        section_title("TECHNICAL SKILLS")
        # Group skills by category (AI/ML, Programming, etc.)
        skills_per_line = 3
        for i in range(0, len(improved["skills"]), skills_per_line):
            skills_group = improved["skills"][i:i+skills_per_line]
            skills_text = "  •  ".join(skills_group)
            for l in wrap_text(skills_text, FONT_BODY, body_size, CONTENT_W):
                items.append((l, FONT_BODY, body_size, 0, 0))

    if improved["experience"]:
        section_title("EXPERIENCE")
        for i, entry in enumerate(improved["experience"]):
            gap = gap_entry if i > 0 else 4
            for l in wrap_text(entry["heading"], FONT_BOLD, body_size + 0.5, CONTENT_W):
                items.append((l, FONT_BOLD, body_size + 0.5, gap, 0))
                gap = 0
            for b in entry["bullets"][:max_bullets_per_entry]:
                for l in wrap_text("• " + b, FONT_BODY, body_size, CONTENT_W - 10):
                    items.append((l, FONT_BODY, body_size, 0, 10))

    if improved["projects"]:
        section_title("PROJECTS")
        for i, entry in enumerate(improved["projects"]):
            gap = gap_entry if i > 0 else 4
            for l in wrap_text(entry["heading"], FONT_BOLD, body_size + 0.5, CONTENT_W):
                items.append((l, FONT_BOLD, body_size + 0.5, gap, 0))
                gap = 0
            for b in entry["bullets"][:max_bullets_per_entry]:
                for l in wrap_text("• " + b, FONT_BODY, body_size, CONTENT_W - 10):
                    items.append((l, FONT_BODY, body_size, 0, 10))

    if improved["education"]:
        section_title("EDUCATION")
        for i, entry in enumerate(improved["education"]):
            gap = gap_entry if i > 0 else 4
            for l in wrap_text(entry["heading"], FONT_BOLD, body_size + 0.5, CONTENT_W):
                items.append((l, FONT_BOLD, body_size + 0.5, gap, 0))
                gap = 0
            for b in entry["bullets"][:max_bullets_per_entry]:
                for l in wrap_text("• " + b, FONT_BODY, body_size, CONTENT_W - 10):
                    items.append((l, FONT_BODY, body_size, 0, 10))

    if improved["certifications"]:
        section_title("CERTIFICATIONS")
        # Break certifications into groups for readability
        certs_per_line = 3
        for i in range(0, len(improved["certifications"]), certs_per_line):
            certs_group = improved["certifications"][i:i+certs_per_line]
            cert_text = "  •  ".join(certs_group)
            for l in wrap_text(cert_text, FONT_BODY, body_size, CONTENT_W):
                items.append((l, FONT_BODY, body_size, 0, 0))

    if improved["achievements"]:
        section_title("ACHIEVEMENTS")
        for a in improved["achievements"]:
            for l in wrap_text("• " + a, FONT_BODY, body_size, CONTENT_W - 10):
                items.append((l, FONT_BODY, body_size, 0, 10))

    for entry in improved["other"]:
        section_title(entry["heading"].upper())
        for b in entry["bullets"][:max_bullets_per_entry]:
            for l in wrap_text("• " + b, FONT_BODY, body_size, CONTENT_W - 10):
                items.append((l, FONT_BODY, body_size, 0, 10))

    return items


def measure_items_height(items):
    total = 0
    for text, font, size, gap, indent in items:
        if text == "__RULE__":
            total += gap
            continue
        total += gap
        total += size * 1.28
    return total


def render_pdf(improved, output_path):
    """Try progressively more compact layouts until the resume fits one page,
    then draw it with reportlab."""
    available_h = PAGE_H - 2 * MARGIN

    configs = [
        dict(body_size=11.0, max_bullets=5, max_summary=3),
        dict(body_size=10.5, max_bullets=5, max_summary=3),
        dict(body_size=10.0, max_bullets=4, max_summary=3),
        dict(body_size=9.5, max_bullets=4, max_summary=2),
        dict(body_size=9.0, max_bullets=4, max_summary=2),
        dict(body_size=8.5, max_bullets=3, max_summary=2),
        dict(body_size=8.0, max_bullets=3, max_summary=1),
    ]

    chosen_items = None
    chosen_size = configs[-1]["body_size"]
    for cfg in configs:
        items = build_pdf_lines(
            improved, cfg["body_size"], cfg["max_bullets"], cfg["max_summary"]
        )
        if measure_items_height(items) <= available_h:
            chosen_items = items
            chosen_size = cfg["body_size"]
            break

    if chosen_items is None:
        cfg = configs[-1]
        chosen_items = build_pdf_lines(
            improved, cfg["body_size"], cfg["max_bullets"], cfg["max_summary"]
        )
        chosen_size = cfg["body_size"]

    c = canvas.Canvas(output_path, pagesize=LETTER)
    y = PAGE_H - MARGIN

    for text, font, size, gap, indent in chosen_items:
        if text == "__RULE__":
            y -= gap
            c.setLineWidth(0.75)
            c.setStrokeColorRGB(0.1, 0.1, 0.1)
            c.line(MARGIN, y, PAGE_W - MARGIN, y)
            y -= 4
            continue
        y -= gap
        line_height = size * 1.28
        y -= line_height * 0.78
        c.setFont(font, size)
        c.setFillColorRGB(0.08, 0.08, 0.08)
        c.drawString(MARGIN + indent, y, text)
        y -= line_height * 0.22

    c.showPage()
    c.save()
    return chosen_size


# --------------------------------------------------------------------------
# Blueprint Routes
# --------------------------------------------------------------------------

@resume_bp.route("/", methods=["GET"])
def resume():
    return render_template("resume.html")


@resume_bp.route("/process", methods=["POST"])
def process():
    jd_text = request.form.get("job_description", "").strip()
    job_role = request.form.get("job_role", "").strip()
    file = request.files.get("resume_file")

    if not jd_text:
        return jsonify({"error": "Please paste the job description."}), 400
    if not file or file.filename == "":
        return jsonify({"error": "Please upload your resume (.pdf or .docx)."}), 400
    if not allowed_file(file.filename):
        return jsonify({"error": "Only .pdf or .docx files are allowed."}), 400

    token = uuid.uuid4().hex[:12]
    ext = file.filename.rsplit(".", 1)[1].lower()
    upload_path = os.path.join(UPLOAD_FOLDER, f"{token}.{ext}")
    file.save(upload_path)

    resume_text = extract_text(upload_path, f"{token}.{ext}")
    if not resume_text.strip():
        return jsonify({"error": "Could not extract text from that file. It may be a scanned image or corrupted."}), 400

    jd_keywords, jd_bigrams = extract_keywords(jd_text)
    result = compute_match(resume_text, jd_keywords, jd_bigrams)

    parsed = parse_resume(resume_text)
    improved = improve_resume(parsed, result["missing"], result["missing_bigrams"], jd_text, job_role)

    output_filename = f"Improved_Resume_{token}.pdf"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)
    render_pdf(improved, output_path)

    # Recompute score against the improved resume text for an "after" preview
    improved_text_bits = improved["skills"] + improved["summary"]
    after_tokens = set(t.lower() for t in improved_text_bits) | set(tokenize(resume_text))
    after_hit = len([k for k in jd_keywords if k in after_tokens or k in {s.lower() for s in improved["skills"]}])
    after_hit += len([b for b in jd_bigrams if b in " ".join(improved["skills"]).lower() or b in resume_text.lower()])
    total = len(jd_keywords) + len(jd_bigrams)
    after_score = round((after_hit / total) * 100, 1) if total else result["score"]
    after_score = max(after_score, result["score"])

    return jsonify({
        "before_score": result["score"],
        "after_score": after_score,
        "matched": result["matched"],
        "missing": result["missing"],
        "matched_bigrams": result["matched_bigrams"],
        "missing_bigrams": result["missing_bigrams"],
        "added_skills": improved.get("_added_skills", []),
        "download_file": output_filename,
    })


@resume_bp.route("/download/<path:filename>")
def download(filename):
    filename = os.path.basename(filename)
    path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(path):
        return jsonify({"error": "File not found. Please run the analysis again."}), 404
    return send_file(
        path,
        as_attachment=True,
        download_name="Improved_Resume.pdf",
        mimetype="application/pdf",
    )